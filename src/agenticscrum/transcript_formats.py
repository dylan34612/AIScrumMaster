"""Convert supported transcript file formats into plain text."""

from __future__ import annotations

import html
import json
import re
from pathlib import Path

_HTML_TAG_RE = re.compile(r"<[^>]+>")
_VTT_VOICE_RE = re.compile(r"<v\s+([^>]+)>(.*?)</v>", re.I)

# Keep cue text bounded so oversized exports stay LLM-friendly.
_MAX_TRANSCRIPT_CHARS = 45_000


def read_transcript_text(path: Path) -> str:
    """Read a transcript file and return normalized plain text."""

    suffix = path.suffix.lower()
    if suffix == ".vtt":
        raw = path.read_text(encoding="utf-8", errors="replace")
        return parse_vtt_to_text(raw)
    if suffix == ".docx":
        return extract_docx_text(path)
    return path.read_text(encoding="utf-8", errors="replace")


def parse_vtt_to_text(vtt: str) -> str:
    """Normalize VTT transcript content into plain speaker/dialogue text."""

    out: list[str] = []
    total_chars = 0
    for raw_line in vtt.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.upper() == "WEBVTT":
            continue
        if "-->" in line:
            continue
        # Numeric cue identifiers
        if line.isdigit():
            continue

        # metadataContent lines: JSON utterances
        if line.startswith("{") and line.endswith("}"):
            try:
                obj = json.loads(line)
            except Exception:
                continue
            spoken = obj.get("spokenText")
            if not isinstance(spoken, str):
                continue
            spoken = " ".join(spoken.strip().split())
            if not spoken:
                continue
            speaker = obj.get("speakerName")
            speaker_text = (
                " ".join(str(speaker).strip().split()) if isinstance(speaker, str) else ""
            )
            rendered = f"{speaker_text}: {spoken}" if speaker_text else spoken
            if total_chars + len(rendered) + 1 > _MAX_TRANSCRIPT_CHARS:
                break
            out.append(rendered)
            total_chars += len(rendered) + 1
            continue

        # content lines: <v Speaker>Name</v>
        match = _VTT_VOICE_RE.match(line)
        if match:
            speaker = " ".join(match.group(1).strip().split())
            spoken_raw = match.group(2).strip()
            spoken = html.unescape(_HTML_TAG_RE.sub(" ", spoken_raw)).strip()
            spoken = " ".join(spoken.split())
            if not spoken:
                continue
            rendered = f"{speaker}: {spoken}" if speaker else spoken
            if total_chars + len(rendered) + 1 > _MAX_TRANSCRIPT_CHARS:
                break
            out.append(rendered)
            total_chars += len(rendered) + 1
            continue

        # Fallback: strip any markup and keep content.
        text = html.unescape(_HTML_TAG_RE.sub(" ", line)).strip()
        text = " ".join(text.split())
        if not text:
            continue
        if total_chars + len(text) + 1 > _MAX_TRANSCRIPT_CHARS:
            break
        out.append(text)
        total_chars += len(text) + 1

    return "\n".join(out).strip()


def extract_docx_text(path: Path) -> str:
    """Extract readable text from a Word .docx transcript."""

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency should be installed
        raise RuntimeError(
            "Reading .docx files requires the python-docx package. "
            "Install project dependencies and try again."
        ) from exc

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
