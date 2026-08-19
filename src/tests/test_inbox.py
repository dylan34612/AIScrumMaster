from __future__ import annotations

from datetime import date
from pathlib import Path

from agenticscrum.config import Settings
from agenticscrum.inbox import (
    finalize_processed_file,
    list_inbox_files,
    parse_inbox_file,
    write_processed_transcript_copy,
)


def _settings(tmp_path: Path, mode: str = "keep", extensions: list[str] | None = None) -> Settings:
    return Settings(
        notes_inbox_enabled=True,
        notes_inbox_path=str(tmp_path),
        notes_inbox_archive_path=str(tmp_path / "_archive"),
        notes_inbox_archive_mode=mode,  # type: ignore[arg-type]
        notes_inbox_min_age_seconds=0,
        notes_inbox_extensions=extensions or [".md", ".txt", ".vtt", ".docx"],
    )


def test_parse_from_filename_date_prefix(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    path = tmp_path / "2026-06-23 - Databricks Stand Up.md"
    path.write_text("These are the notes.", encoding="utf-8")
    note = parse_inbox_file(settings, path)
    assert note.title == "Databricks Stand Up"
    assert note.meeting_date == date(2026, 6, 23)
    assert note.notes == "These are the notes."


def test_parse_from_filename_date_suffix(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    path = tmp_path / "Weekly Working Session - 2026-06-22.txt"
    path.write_text("Notes body.", encoding="utf-8")
    note = parse_inbox_file(settings, path)
    assert note.title == "Weekly Working Session"
    assert note.meeting_date == date(2026, 6, 22)
    assert note.notes == "Notes body."


def test_parse_yaml_front_matter(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    path = tmp_path / "anything.md"
    path.write_text(
        """---
title: Demo Meeting
date: 2026-06-21
---

Line 1
Line 2
""",
        encoding="utf-8",
    )
    note = parse_inbox_file(settings, path)
    assert note.title == "Demo Meeting"
    assert note.meeting_date == date(2026, 6, 21)
    assert note.notes == "Line 1\nLine 2"


def test_parse_key_value_header_block(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    path = tmp_path / "random.txt"
    path.write_text(
        """Meeting Title: Standup
Meeting Date: 6/23/2026

Did X
Did Y
""",
        encoding="utf-8",
    )
    note = parse_inbox_file(settings, path)
    assert note.title == "Standup"
    assert note.meeting_date == date(2026, 6, 23)
    assert note.notes == "Did X\nDid Y"


def test_parse_markdown_h1(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    path = tmp_path / "title-only.md"
    path.write_text(
        """# My Meeting

Body line
""",
        encoding="utf-8",
    )
    note = parse_inbox_file(settings, path)
    assert note.title == "My Meeting"
    assert note.notes == "Body line"


def test_list_inbox_files_filters_by_extension(tmp_path: Path) -> None:
    settings = _settings(tmp_path, extensions=[".md", ".txt", ".vtt", ".docx"])
    (tmp_path / "a.md").write_text("x", encoding="utf-8")
    (tmp_path / "b.txt").write_text("y", encoding="utf-8")
    (tmp_path / "c.vtt").write_text("WEBVTT\n\n00:00:00.000 --> 00:00:01.000\nHello\n", encoding="utf-8")
    (tmp_path / "d.docx").write_bytes(b"PK")  # placeholder; only filtered by extension here
    (tmp_path / "e.pdf").write_text("z", encoding="utf-8")
    paths = list_inbox_files(settings)
    assert [p.name for p in paths] == ["a.md", "b.txt", "c.vtt", "d.docx"]


def test_parse_vtt_inbox_file(tmp_path: Path) -> None:
    settings = _settings(tmp_path)
    path = tmp_path / "2026-06-23 - Standup.vtt"
    path.write_text(
        """WEBVTT

1
00:00:00.000 --> 00:00:02.000
<v Alice>Ship the inbox formats today</v>

2
00:00:02.500 --> 00:00:04.000
<v Bob>Sounds good</v>
""",
        encoding="utf-8",
    )
    note = parse_inbox_file(settings, path)
    assert note.title == "Standup"
    assert note.meeting_date == date(2026, 6, 23)
    assert note.notes == "Alice: Ship the inbox formats today\nBob: Sounds good"


def test_parse_docx_inbox_file(tmp_path: Path) -> None:
    from docx import Document

    settings = _settings(tmp_path)
    path = tmp_path / "2026-06-23 - Planning.docx"
    doc = Document()
    doc.add_paragraph("Meeting Title: Planning")
    doc.add_paragraph("Meeting Date: 2026-06-23")
    doc.add_paragraph("")
    doc.add_paragraph("Discussed backlog grooming.")
    doc.add_paragraph("Assigned follow-ups.")
    doc.save(path)

    note = parse_inbox_file(settings, path)
    assert note.title == "Planning"
    assert note.meeting_date == date(2026, 6, 23)
    assert "Discussed backlog grooming." in note.notes
    assert "Assigned follow-ups." in note.notes


def test_finalize_processed_file_archives(tmp_path: Path) -> None:
    settings = _settings(tmp_path, mode="archive")
    path = tmp_path / "2026-06-23 - Demo.md"
    path.write_text("Body", encoding="utf-8")
    note = parse_inbox_file(settings, path)
    dest = finalize_processed_file(settings, note)
    assert dest is not None
    assert dest.exists()
    assert not path.exists()


def test_write_processed_transcript_copy(tmp_path: Path) -> None:
    settings = _settings(tmp_path, mode="keep")
    dest = write_processed_transcript_copy(
        settings,
        title="Demo Meeting",
        meeting_date=date(2026, 6, 23),
        notes="Line 1\nLine 2",
        source="Manual UI",
        run_id=123,
        run_status="Success",
    )
    assert dest.exists()
    text = dest.read_text(encoding="utf-8")
    assert "Meeting Title: Demo Meeting" in text
    assert "Meeting Date: 2026-06-23" in text
    assert "Ingestion Run ID: 123" in text
    assert "Line 1" in text

