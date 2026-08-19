"""Tests for transcript format converters."""

from __future__ import annotations

from pathlib import Path

from agenticscrum.transcript_formats import extract_docx_text, parse_vtt_to_text, read_transcript_text


def test_parse_vtt_voice_tags() -> None:
    text = parse_vtt_to_text(
        """WEBVTT

00:00:00.000 --> 00:00:01.500
<v Dana>Hello team</v>

00:00:01.500 --> 00:00:03.000
<v Sam>Ready to start</v>
"""
    )
    assert text == "Dana: Hello team\nSam: Ready to start"


def test_parse_vtt_json_utterances() -> None:
    text = parse_vtt_to_text(
        """WEBVTT

{"speakerName":"Alex","spokenText":"First point"}
{"speakerName":"Blair","spokenText":"Second point"}
"""
    )
    assert text == "Alex: First point\nBlair: Second point"


def test_read_transcript_text_docx(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "notes.docx"
    doc = Document()
    doc.add_paragraph("Line one")
    doc.add_paragraph("Line two")
    doc.save(path)

    assert read_transcript_text(path) == "Line one\nLine two"
    assert extract_docx_text(path) == "Line one\nLine two"
